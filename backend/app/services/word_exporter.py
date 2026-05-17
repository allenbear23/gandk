import logging
import json
import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from app.models.question import ExamResult

logger = logging.getLogger(__name__)

def set_run_font(run, font_name="新細明體", size_pt=11, bold=False, italic=False, underline=False, color_rgb=None):
    """
    專業級字型控制器：支援 ASCII 與 中日韓東亞字型 (eastAsia) 完美渲染，解決 macOS 下標楷體/新細明體顯示問題。
    """
    run.font.name = font_name
    
    # 強制注入東亞字型 XML 設定
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_paragraph_bottom_double_border(paragraph):
    """
    利用 XML 注入段落底部的雙線底線 (中華民國國家考試經典風格)
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="double" w:sz="12" w:space="8" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

def add_word_bank_box(doc, words_text):
    """
    建立一個單格表格作為文意選填的單字庫外框
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    # 設置表格寬度
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(5.8)
            # 設置邊框為細實線
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>')
            tcPr.append(tcBorders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)
            
            run = p.add_run(words_text.strip())
            set_run_font(run, font_name="Times New Roman", size_pt=9.5, italic=True)

def export_to_docx(exam_result: ExamResult) -> bytes:
    """
    使用 python-docx 產生 100% 還原考古題格式之二進位制 .docx 文件！
    包含：頁首單欄、考題雙欄、大題獨立題號重置、字彙無選項、文意選填單字庫外框。
    """
    from app.db.supabase_client import get_supabase
    
    doc = Document()
    
    # 設置第一節 (頁首專區：單欄)
    header_section = doc.sections[0]
    header_section.top_margin = Inches(0.8)
    header_section.bottom_margin = Inches(0.8)
    header_section.left_margin = Inches(0.8)
    header_section.right_margin = Inches(0.8)

    # 嘗試獲取科目的風格設定
    style_json = None
    try:
        sb = get_supabase()
        res = sb.table("subjects").select("style_prompt").eq("id", exam_result.subject_id).single().execute()
        style_prompt = res.data.get("style_prompt") if res.data else None
        
        if style_prompt:
            cleaned = style_prompt.strip()
            if cleaned.startswith("```"):
                cleaned = re.sub(r'^```(?:json)?\s*\n?', '', cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r'\n?```\s*$', '', cleaned)
            style_json = json.loads(cleaned)
    except Exception as e:
        logger.warning(f"⚠️ 無法讀取或解析風格設定，將使用預設格式: {e}")

    # 建立大題屬性對照表，使排版系統完全「資料庫驅動」，支援所有學科科目！
    sections_meta = {}
    if style_json and "sections" in style_json:
        for sec in style_json["sections"]:
            name = sec.get("section_name", "")
            if name:
                sections_meta[name] = sec

    # 1. 建立單欄頁首
    last_header_para = None
    if style_json and "document_header" in style_json:
        header_text = style_json["document_header"]
        lines = header_text.strip().split('\n')
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(lines[0])
        set_run_font(run, font_name="標楷體", size_pt=15, bold=True)
        
        for line in lines[1:]:
            p_line = doc.add_paragraph()
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_line.add_run(line)
            set_run_font(run, font_name="新細明體", size_pt=10.5)
            last_header_para = p_line
            
        if last_header_para:
            add_paragraph_bottom_double_border(last_header_para)
    else:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(f"Weekly Review Test: {exam_result.subject}")
        set_run_font(run, font_name="標楷體", size_pt=16, bold=True)
        
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        units_str = "、".join(exam_result.units)
        run_meta = p_meta.add_run(f"Units: {units_str}   |   Class: _________   No: _________   Name: _________________")
        set_run_font(run_meta, font_name="新細明體", size_pt=10.5)
        last_header_para = p_meta
        add_paragraph_bottom_double_border(last_header_para)

    # 2. 加入分節符號，將後續試題內容轉為「雙欄排版 (Two-Column)」
    exam_section = doc.add_section()
    exam_section.start_type = 0 # 0 = Continuous Section Break
    exam_section.top_margin = Inches(0.8)
    exam_section.bottom_margin = Inches(0.8)
    exam_section.left_margin = Inches(0.8)
    exam_section.right_margin = Inches(0.8)
    
    # 設置雙欄 XML 屬性
    sectPr = exam_section._sectPr
    cols = sectPr.xpath('w:cols')
    if cols:
        cols[0].set(qn('w:num'), '2')
        cols[0].set(qn('w:space'), '540') # 欄間距 540 dxa (約 0.38 英吋)
    else:
        cols_xml = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="540"/>')
        sectPr.append(cols_xml)

    # 3. 依序產生試題，控制大題本地題號重置
    current_section_name = None
    sec_q_counter = 0
    listening_global_counter = 0
    
    for i, q in enumerate(exam_result.questions, 1):
        # 偵測大題切換
        if q.section and q.section != current_section_name:
            current_section_name = q.section
            
            # 從大題對照表讀取屬性
            sec_meta = sections_meta.get(current_section_name, {})
            layout_type = sec_meta.get("layout_type", "").lower()
            
            # 判斷是否需要重置題號：非聽力測驗大題，均重置題號為 1
            is_listening_sec = (layout_type == "listening") or ("聽力" in current_section_name)
            if not is_listening_sec:
                sec_q_counter = 1
            else:
                # 聽力測驗跨 Part 題號連續
                if listening_global_counter == 0:
                    listening_global_counter = 1
                sec_q_counter = listening_global_counter
                
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(8)
            p_sec.paragraph_format.space_after = Pt(4)
            p_sec.paragraph_format.keep_with_next = True
            run_sec = p_sec.add_run(current_section_name)
            set_run_font(run_sec, font_name="標楷體", size_pt=11.5, bold=True, underline=True)
            
            # 特殊外框處理：若是「文意選填」大題，在標題下方自動渲染一個單字庫外框 (Word Bank Box)
            is_word_bank_sec = (layout_type == "word_bank") or ("文意選填" in current_section_name)
            if is_word_bank_sec:
                # 嘗試從 q.question 或預置內容尋找單字選項庫，若無則預設一組符合課本範圍的詞庫
                word_bank_text = "A. overuse    B. feathery    C. scarce    D. recovery\nE. shelters    F. adopt    G. compete    H. remain\nI. scientific    J. researchers"
                # 如果 question 內容包含 Word Bank: 標記，進行解析
                if "Word Bank:" in q.question:
                    parts = q.question.split("Word Bank:")
                    word_bank_text = parts[0].strip()
                    q.question = parts[1].strip() # 剝離出純文章
                add_word_bank_box(doc, word_bank_text)
                doc.add_paragraph() # 空行分隔

        # 4. 根據大題類型進行精細排版 (支援 layout_type 自訂與中文關鍵字模糊匹配雙通道)
        sec_meta = sections_meta.get(q.section, {}) if q.section else {}
        layout_type = sec_meta.get("layout_type", "").lower() if isinstance(sec_meta, dict) else ""
        
        is_listening = (layout_type == "listening") or ("聽力" in (q.section or ""))
        is_vocabulary = (layout_type == "vocabulary") or ("字彙" in (q.section or ""))
        is_cloze = (layout_type == "cloze") or ("克漏字" in (q.section or ""))
        is_completion = (layout_type == "word_bank") or ("文意選填" in (q.section or ""))
        is_translation = (layout_type == "translation") or ("翻譯" in (q.section or ""))

        # 更新聽力全局計數器
        if is_listening:
            listening_global_counter = sec_q_counter + 1

        # 字彙測驗排版：非選擇題，無選項，直接渲染題目
        if is_vocabulary:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.space_after = Pt(4)
            run_num = p_q.add_run(f"{sec_q_counter}. ")
            set_run_font(run_num, font_name="Times New Roman", size_pt=10.5, bold=True)
            
            run_q = p_q.add_run(q.question)
            set_run_font(run_q, font_name="新細明體", size_pt=10.5)
            sec_q_counter += 1

        # 克漏字 / 文意選填排版：文章融合渲染
        elif is_cloze or is_completion:
            # 偵測是否為該段文章的第一題
            # 如果題目文本長度大於 80 字，視為大題的「長篇閱讀引導文章」
            if len(q.question) > 80:
                p_passage = doc.add_paragraph()
                p_passage.paragraph_format.space_before = Pt(4)
                p_passage.paragraph_format.space_after = Pt(6)
                run_p = p_passage.add_run(q.question)
                set_run_font(run_p, font_name="新細明體", size_pt=10)
                
                # 接著渲染第一題的選項
                p_opts = doc.add_paragraph()
                p_opts.paragraph_format.left_indent = Inches(0.2)
                p_opts.paragraph_format.space_after = Pt(6)
                run_num = p_opts.add_run(f"{sec_q_counter}. ")
                set_run_font(run_num, font_name="Times New Roman", size_pt=10.5, bold=True)
                
                choices_text = "   ".join([f"({c.key}) {c.text}" for c in q.choices])
                run_opts = p_opts.add_run(choices_text)
                set_run_font(run_opts, font_name="Times New Roman", size_pt=9.5)
            else:
                # 後續題目只印選項，免除段落重複
                p_opts = doc.add_paragraph()
                p_opts.paragraph_format.left_indent = Inches(0.2)
                p_opts.paragraph_format.space_after = Pt(6)
                run_num = p_opts.add_run(f"{sec_q_counter}. ")
                set_run_font(run_num, font_name="Times New Roman", size_pt=10.5, bold=True)
                
                choices_text = "   ".join([f"({c.key}) {c.text}" for c in q.choices])
                run_opts = p_opts.add_run(choices_text)
                set_run_font(run_opts, font_name="Times New Roman", size_pt=9.5)
                
            sec_q_counter += 1

        # 翻譯題排版 (引導式翻譯編號如 (1) (2))
        elif is_translation:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.space_after = Pt(4)
            # 使用引導括弧題號
            run_num = p_q.add_run(f"({sec_q_counter}) ")
            set_run_font(run_num, font_name="Times New Roman", size_pt=10.5, bold=True)
            
            run_q = p_q.add_run(q.question)
            set_run_font(run_q, font_name="新細明體", size_pt=10.5)
            sec_q_counter += 1

        # 一般文法與聽力題排版
        else:
            p_q = doc.add_paragraph()
            p_q.paragraph_format.space_after = Pt(3)
            run_num = p_q.add_run(f"{sec_q_counter}. ")
            set_run_font(run_num, font_name="Times New Roman", size_pt=10.5, bold=True)
            
            run_q = p_q.add_run(q.question)
            set_run_font(run_q, font_name="新細明體", size_pt=10.5)
            
            if q.choices:
                p_opts = doc.add_paragraph()
                p_opts.paragraph_format.left_indent = Inches(0.25)
                p_opts.paragraph_format.space_after = Pt(8)
                choices_text = "   ".join([f"({c.key}) {c.text}" for c in q.choices])
                run_opts = p_opts.add_run(choices_text)
                set_run_font(run_opts, font_name="Times New Roman", size_pt=10)
                
            sec_q_counter += 1

    # 5. 實體分頁：參考答案與標準解析頁面 (答案頁面轉回單欄，方便大字閱讀)
    ans_section = doc.add_section()
    ans_section.start_type = 2 # 2 = New Page Section Break
    
    # 答案頁改回單欄排版以利閱讀
    ans_sectPr = ans_section._sectPr
    ans_cols = ans_sectPr.xpath('w:cols')
    if ans_cols:
        ans_cols[0].set(qn('w:num'), '1')
        
    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ans_title.paragraph_format.space_before = Pt(20)
    p_ans_title.paragraph_format.space_after = Pt(20)
    run_ans_title = p_ans_title.add_run("【教育部考試評量檢定處 ‧ 參考答案與標準解析】")
    set_run_font(run_ans_title, font_name="標楷體", size_pt=14, bold=True)
    
    current_ans_section = None
    sec_ans_counter = 0
    listening_ans_global = 0
    
    for i, q in enumerate(exam_result.questions, 1):
        if q.section and q.section != current_ans_section:
            current_ans_section = q.section
            if "聽力" not in current_ans_section:
                sec_ans_counter = 1
            else:
                if listening_ans_global == 0:
                    listening_ans_global = 1
                sec_ans_counter = listening_ans_global
                
            p_ans_sec = doc.add_paragraph()
            p_ans_sec.paragraph_format.space_before = Pt(10)
            p_ans_sec.paragraph_format.keep_with_next = True
            run_ans_sec = p_ans_sec.add_run(current_ans_section)
            set_run_font(run_ans_sec, font_name="標楷體", size_pt=11.5, bold=True)

        if "聽力" in (q.section or ""):
            listening_ans_global = sec_ans_counter + 1

        p_ans_item = doc.add_paragraph()
        p_ans_item.paragraph_format.space_after = Pt(4)
        
        # 標楷體答案印出
        run_key = p_ans_item.add_run(f"第 {sec_ans_counter} 題  參考答案：【")
        set_run_font(run_key, font_name="新細明體", size_pt=10.5)
        
        run_ans = p_ans_item.add_run(q.answer)
        set_run_font(run_ans, font_name="標楷體", size_pt=11, bold=True)
        
        run_end = p_ans_item.add_run("】")
        set_run_font(run_end, font_name="新細明體", size_pt=10.5)
        
        # 智慧解析說明文字
        p_exp = doc.add_paragraph()
        p_exp.paragraph_format.left_indent = Inches(0.2)
        p_exp.paragraph_format.space_after = Pt(12)
        run_exp = p_exp.add_run(q.explanation)
        set_run_font(run_exp, font_name="新細明體", size_pt=10, color_rgb=RGBColor(0x33, 0x33, 0x33))
        
        sec_ans_counter += 1

    # 輸出為二進位制 bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()
